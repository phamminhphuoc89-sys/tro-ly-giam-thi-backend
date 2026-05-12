# app/utils/legacy/bao_cao_thi_dua.py
# -*- coding: utf-8 -*-
"""
Logic xử lý báo cáo thi đua - không phụ thuộc Tkinter
Các hàm: đọc file Excel điểm SĐB, lỗi tập thể, lỗi cá nhân,
tạo comment, cập nhật mẫu Word.
"""

import re
import unicodedata
from datetime import date, timedelta
from collections import Counter
from typing import Dict, List, Any
import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------- Bảng tháng tiếng Anh ----------
month_names_en_full = {
    1: "January", 2: "February", 3: "March", 4: "April",
    5: "May", 6: "June", 7: "July", 8: "August",
    9: "September", 10: "October", 11: "November", 12: "December"
}

# ---------- Hàm tiện ích ----------
def remove_accents(text: str) -> str:
    text = unicodedata.normalize('NFKD', text).encode('ASCII', 'ignore').decode('ASCII')
    return text.lower()

def get_week_info(d: date):
    monday = d - timedelta(days=d.weekday())
    month = monday.month
    year = monday.year
    first_day = date(year, month, 1)
    days_to_first_monday = (7 - first_day.weekday()) % 7
    first_monday = first_day + timedelta(days=days_to_first_monday)
    week_num = (monday - first_monday).days // 7 + 1
    if week_num < 1:
        week_num = 1
    return week_num, month, year

def make_ngay_thang_vi(start_date: date, end_date: date) -> str:
    week_num, month, year = get_week_info(start_date)
    return f"Tuần {week_num} -- Tháng {month:02d}/{year} ({start_date.strftime('%d/%m/%Y')} -- {end_date.strftime('%d/%m/%Y')})"

def make_ngay_thang_en(start_date: date, end_date: date) -> str:
    return f"from {start_date.strftime('%d')} {month_names_en_full[start_date.month]} {start_date.year} to {end_date.strftime('%d')} {month_names_en_full[end_date.month]} {end_date.year}"

# ---------- Định dạng bảng Word ----------
def set_cell_border(cell, border_size=4, border_color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ['top', 'left', 'bottom', 'right']:
        tag = f'w:{edge}'
        border = OxmlElement(tag)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)
        tcPr.append(border)

def set_table_borders(table, border_size=4, border_color="000000"):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, border_size, border_color)

def format_cell_font(cell, size=10, font_name="Times New Roman", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = font_name
            run.font.bold = bold
        if not paragraph.runs and paragraph.text:
            run = paragraph.add_run(paragraph.text)
            run.font.size = Pt(size)
            run.font.name = font_name
            run.font.bold = bold
            paragraph.clear()
            paragraph.add_run(run.text)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def format_table_cells(table, size=10, font_name="Times New Roman",
                       header_bold=True, align_center_columns=None):
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            align = WD_ALIGN_PARAGRAPH.CENTER if (align_center_columns and col_idx in align_center_columns) else WD_ALIGN_PARAGRAPH.LEFT
            is_header = (row_idx == 0)
            bold = header_bold and is_header
            format_cell_font(cell, size, font_name, bold, align)
    set_table_borders(table)

# ---------- Đọc dữ liệu Excel ----------
def read_diem_sdb(filepath: str) -> Dict[str, Any]:
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    ws = wb["Diem SDB"]
    rows = list(ws.iter_rows(values_only=True))
    headers = rows[0]
    week_cols = {}
    for i, h in enumerate(headers):
        if h and "Điểm SĐB Tuần" in str(h):
            try:
                week_cols[int(str(h).replace("Điểm SĐB Tuần", "").strip())] = i
            except ValueError:
                pass
    data_rows = [r for r in rows[1:] if r[0] is not None and str(r[0]).strip()]
    latest = 1
    for wk in sorted(week_cols, reverse=True):
        if any(r[week_cols[wk]] is not None for r in data_rows):
            latest = wk
            break
    col = week_cols.get(latest, 2)
    classes = [
        {"lop": str(r[0]).strip(), "gvcn": str(r[1]).strip() if r[1] else "",
         "diem": float(r[col]) if r[col] is not None else 0.0}
        for r in data_rows if r[0]
    ]
    wb.close()
    return {"week": latest, "week_label": f"Tuần {latest}", "classes": classes}

def read_loi_tap_the(filepath: str) -> Dict[str, Any]:
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheet = next((s for s in wb.sheetnames if s.lower().startswith("tuan")), None)
    if not sheet:
        wb.close()
        return {"sheet_name": "", "records": []}
    ws = wb[sheet]
    rows = list(ws.iter_rows(values_only=True))
    hi = next((i for i, r in enumerate(rows) if r[0] == "Ngày tháng"), None)
    if hi is None:
        wb.close()
        return {"sheet_name": sheet, "records": []}
    def fmt(v):
        return v.strftime("%d-%m-%Y") if hasattr(v, "strftime") else str(v)[:10]
    records = [
        {"ngay": fmt(r[0]), "tiet": str(r[1]) if r[1] is not None else "",
         "lop": str(r[2]) if r[2] else "", "mon": str(r[3]) if r[3] else "",
         "nguon": str(r[4]) if r[4] else "", "noi_dung": str(r[5]) if r[5] else "",
         "ma_loi": str(r[6]) if r[6] is not None else "",
         "diem_tru": str(r[7]) if r[7] is not None else ""}
        for r in rows[hi+1:] if r[0]
    ]
    wb.close()
    return {"sheet_name": sheet, "records": records}

def read_loi_ca_nhan(filepath: str) -> List[Dict]:
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    ws = wb["Tong hop"]
    rows = list(ws.iter_rows(values_only=True))
    def fmt(v):
        return v.strftime("%d-%m-%Y") if hasattr(v, "strftime") else str(v)[:10] if v else ""
    records = [
        {"ngay": fmt(r[0]), "tiet": str(r[1]) if r[1] is not None else "",
         "lop": str(r[2]) if r[2] else "", "mon": str(r[3]) if r[3] else "",
         "nguon": str(r[4]) if r[4] else "", "ho_ten": str(r[5]) if r[5] else "",
         "noi_dung": str(r[6]) if r[6] else "", "cap_do": str(r[7]) if r[7] is not None else "",
         "diem": str(r[8]) if r[8] is not None else ""}
        for r in rows[1:] if r[5]
    ]
    wb.close()
    return records

# ---------- Nhận xét ----------
def generate_comments(diem_sdb, loi_tap_the, loi_ca_nhan):
    sorted_cls = sorted(diem_sdb["classes"], key=lambda x: x["diem"], reverse=True)
    top_class = sorted_cls[0]["lop"] if sorted_cls else ""
    bottom_class = sorted_cls[-1]["lop"] if sorted_cls else ""

    tt_class_cnt = Counter(r["lop"] for r in loi_tap_the["records"])
    worst_tt_class = tt_class_cnt.most_common(1)[0][0] if tt_class_cnt else ""
    cn_class_cnt = Counter(r["lop"] for r in loi_ca_nhan)
    worst_cn_class = cn_class_cnt.most_common(1)[0][0] if cn_class_cnt else ""

    tinh_hinh_vi = []
    if top_class:
        tinh_hinh_vi.append(f"Lớp {top_class} dẫn đầu về điểm Sổ đầu bài trong tuần.")
    if loi_tap_the["records"]:
        tinh_hinh_vi.append(f"Ghi nhận {len(loi_tap_the['records'])} vi phạm tập thể, tập trung nhiều ở lớp {worst_tt_class}.")
    if loi_ca_nhan:
        tinh_hinh_vi.append(f"Có {len(loi_ca_nhan)} vi phạm cá nhân, nhiều nhất ở lớp {worst_cn_class}.")
        err_list = [r["noi_dung"] for r in loi_ca_nhan if r["noi_dung"]]
        common_err = Counter(err_list).most_common(3)
        if common_err:
            tinh_hinh_vi.append(f"Các lỗi thường gặp: {', '.join([e for e, _ in common_err])}.")
    else:
        tinh_hinh_vi.append("Không có vi phạm cá nhân nào trong tuần.")

    tinh_hinh_en = []
    if top_class:
        tinh_hinh_en.append(f"Class {top_class} achieved the highest Register Notebook score.")
    if loi_tap_the["records"]:
        tinh_hinh_en.append(f"{len(loi_tap_the['records'])} class violations recorded, mostly in class {worst_tt_class}.")
    if loi_ca_nhan:
        tinh_hinh_en.append(f"{len(loi_ca_nhan)} individual violations, most in class {worst_cn_class}.")
    else:
        tinh_hinh_en.append("No individual violations this week.")

    tuyen_duong_vi = f"Tuyên dương tập thể lớp {top_class} đã đạt điểm Sổ đầu bài cao nhất khối, thể hiện sự cố gắng trong duy trì nề nếp học tập." if top_class else "Chưa có lớp nào đạt thành tích xuất sắc, cần cố gắng hơn."
    phe_binh_vi = f"Cần chấn chỉnh lớp {bottom_class} do điểm Sổ đầu bài thấp nhất khối." if bottom_class else ""
    if loi_tap_the["records"] or loi_ca_nhan:
        phe_binh_vi += f" Đặc biệt lớp {worst_tt_class or worst_cn_class} có nhiều vi phạm, đề nghị nghiêm túc rút kinh nghiệm."

    commendation_en = f"Commendation is given to Class {top_class} for achieving the highest Register Notebook score." if top_class else "No outstanding class this week."
    criticism_en = f"Class {bottom_class} needs improvement in Register Notebook scores." if bottom_class else "Overall discipline is acceptable, keep maintaining."

    return {
        "tinh_hinh_chung_vi": "\n".join(f"- {line}" for line in tinh_hinh_vi),
        "tinh_hinh_chung_en": "\n".join(f"- {line}" for line in tinh_hinh_en),
        "tuyen_duong": tuyen_duong_vi,
        "phe_binh": phe_binh_vi,
        "commendation": commendation_en,
        "criticism": criticism_en
    }

# ---------- Cập nhật template (không cần start_date, end_date toàn cục) ----------
def update_template(template_path, diem_sdb, loi_tap_the, loi_ca_nhan, comments, start_date, end_date, output_path):
    doc = Document(template_path)

    # Cập nhật bảng điểm SĐB
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = ' '.join(cell.text for cell in table.rows[0].cells).lower()
            if 'lớp' in header_text and 'giáo viên chủ nhiệm' in header_text:
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)
                for cls in diem_sdb["classes"]:
                    row = table.add_row()
                    row.cells[0].text = cls["lop"]
                    row.cells[1].text = cls["gvcn"]
                    row.cells[2].text = f"{cls['diem']:.1f}"
                format_table_cells(table, size=10, font_name="Times New Roman", header_bold=True, align_center_columns=[0,2])
                break

    # Bảng lỗi tập thể
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = ' '.join(cell.text for cell in table.rows[0].cells).lower()
            if 'ngày tháng' in header_text and 'tiết' in header_text and 'lớp' in header_text:
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)
                for r in loi_tap_the["records"]:
                    row = table.add_row()
                    row.cells[0].text = r["ngay"]
                    row.cells[1].text = r["tiet"]
                    row.cells[2].text = r["lop"]
                    row.cells[3].text = r["mon"]
                    row.cells[4].text = r["nguon"]
                    row.cells[5].text = r["noi_dung"]
                    row.cells[6].text = r["ma_loi"]
                    row.cells[7].text = r["diem_tru"]
                format_table_cells(table, size=10, font_name="Times New Roman", header_bold=True, align_center_columns=[0,1,2,3,4,6,7])
                break

    # Bảng lỗi cá nhân
    for table in doc.tables:
        if len(table.rows) > 0:
            header_text = ' '.join(cell.text for cell in table.rows[0].cells).lower()
            if 'ngày tháng' in header_text and 'họ và tên' in header_text:
                while len(table.rows) > 1:
                    table._element.remove(table.rows[-1]._element)
                for r in loi_ca_nhan:
                    row = table.add_row()
                    row.cells[0].text = r["ngay"]
                    row.cells[1].text = r["tiet"]
                    row.cells[2].text = r["lop"]
                    row.cells[3].text = r["mon"]
                    row.cells[4].text = r["nguon"]
                    row.cells[5].text = r["ho_ten"]
                    row.cells[6].text = r["noi_dung"]
                    row.cells[7].text = r["cap_do"]
                    row.cells[8].text = r["diem"]
                format_table_cells(table, size=10, font_name="Times New Roman", header_bold=True, align_center_columns=[0,1,2,3,4,7,8])
                break

    # --- THAY THẾ THÔNG MINH (chỉ thay số, giữ nguyên định dạng) ---
    def smart_replace_paragraph(para, s_date, e_date):
        text = para.text
        new_text = text

        # Tiếng Việt
        if re.search(r'Tuần\s+\d+', text) and re.search(r'Tháng\s+\d{1,2}/\d{4}', text):
            week_num, month, year = get_week_info(s_date)
            start_str = s_date.strftime('%d/%m/%Y')
            end_str = e_date.strftime('%d/%m/%Y')
            new_text = re.sub(r'(Tuần\s+)\d+', r'\g<1>' + str(week_num), new_text)
            new_text = re.sub(r'(Tháng\s+)\d{1,2}/\d{4}', r'\g<1>' + f"{month}/{year}", new_text)
            occurrence = [0]
            def _replace_date(m):
                occurrence[0] += 1
                return start_str if occurrence[0] == 1 else end_str
            new_text = re.sub(r'\d{2}/\d{2}/\d{4}', _replace_date, new_text)

        # Tiếng Anh
        if re.search(r'\bfrom\b', new_text, re.IGNORECASE) and re.search(r'\bto\b', new_text, re.IGNORECASE):
            d_start = s_date.strftime('%d')
            m_start = month_names_en_full[s_date.month]
            d_end   = e_date.strftime('%d')
            m_end   = month_names_en_full[e_date.month]
            yr      = str(e_date.year)
            pat_a = r'from\s+\d{1,2}\s+[A-Za-z]+\s+to\s+\d{1,2}\s+[A-Za-z]+\s+\d{4}'
            if re.search(pat_a, new_text, re.IGNORECASE):
                new_text = re.sub(pat_a,
                                  f'from {d_start} {m_start} to {d_end} {m_end} {yr}',
                                  new_text, flags=re.IGNORECASE)
            else:
                pat_b = r'from\s+\d{1,2}\s+[A-Za-z]+\s+\d{4}\s+to\s+\d{1,2}\s+[A-Za-z]+\s+\d{4}'
                new_text = re.sub(pat_b,
                                  f'from {d_start} {m_start} {s_date.year} to {d_end} {m_end} {yr}',
                                  new_text, flags=re.IGNORECASE)

        if new_text == text:
            return False

        # Giữ nguyên định dạng
        dominant = {}
        for run in para.runs:
            if run.text.strip():
                dominant = {
                    'bold':      run.bold,
                    'italic':    run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                }
                try:
                    rPr = run._element.find(qn('w:rPr'))
                    color_el = rPr.find(qn('w:color')) if rPr is not None else None
                    val = color_el.get(qn('w:val')) if color_el is not None else None
                    dominant['color_val'] = val if val and val != 'auto' else None
                except Exception:
                    dominant['color_val'] = None
                break

        p_elem = para._element
        pPr = p_elem.find(qn('w:pPr'))
        pPr_copy = copy.deepcopy(pPr) if pPr is not None else None
        para.clear()
        if pPr_copy is not None:
            p_elem.insert(0, pPr_copy)

        new_run = para.add_run(new_text)
        new_run.bold      = dominant.get('bold')
        new_run.italic    = dominant.get('italic')
        new_run.underline = dominant.get('underline')
        if dominant.get('font_name'):
            new_run.font.name = dominant['font_name']
        if dominant.get('font_size'):
            new_run.font.size = dominant['font_size']
        if dominant.get('color_val'):
            from docx.shared import RGBColor
            try:
                hex_val = dominant['color_val']
                r_val = int(hex_val[0:2], 16)
                g_val = int(hex_val[2:4], 16)
                b_val = int(hex_val[4:6], 16)
                new_run.font.color.rgb = RGBColor(r_val, g_val, b_val)
            except Exception:
                pass
        return True

    for para in doc.paragraphs:
        smart_replace_paragraph(para, start_date, end_date)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    smart_replace_paragraph(para, start_date, end_date)

    # Thay thế placeholder nhận xét
    replacements = {
        "tinh_hinh_chung_vi": comments["tinh_hinh_chung_vi"],
        "tinh_hinh_chung_en": comments["tinh_hinh_chung_en"],
        "tuyen_duong": comments["tuyen_duong"],
        "phe_binh": comments["phe_binh"],
        "commendation": comments["commendation"],
        "criticism": comments["criticism"]
    }
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if f"{{{{{key}}}}}" in para.text:
                new_text = para.text.replace(f"{{{{{key}}}}}", value)
                para.clear()
                run = para.add_run(new_text)
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for key, value in replacements.items():
                        if f"{{{{{key}}}}}" in para.text:
                            new_text = para.text.replace(f"{{{{{key}}}}}", value)
                            para.clear()
                            run = para.add_run(new_text)
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(11)

    doc.save(output_path)